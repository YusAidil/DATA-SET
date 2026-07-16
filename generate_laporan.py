# -*- coding: utf-8 -*-
"""Membuat file laporan Word (Laporan_Praktikum4.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def judul(doc, teks, level):
    return doc.add_heading(teks, level=level)


def par(doc, teks):
    p = doc.add_paragraph(teks)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def sisip(doc, teks):
    p = doc.add_paragraph()
    r = p.add_run(f"[ {teks} ]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


doc = Document()

# ------- Judul -------
t = doc.add_heading("PRAKTIKUM IV", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("UNSUPERVISED LEARNING : K-MEANS")
r.bold = True
r.font.size = Pt(14)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run("Klasterisasi Tingkat Kerentanan Sosial-Ekonomi Kabupaten/Kota di Indonesia")
r2.bold = True

# ============================================================
judul(doc, "1. MATERI SINGKAT", 1)

judul(doc, "A. Unsupervised Learning", 2)
par(doc, "Unsupervised learning merupakan salah satu pendekatan dalam machine learning yang "
    "digunakan untuk mengolah data yang belum memiliki label atau kategori. Pada metode ini, "
    "komputer tidak diberi tahu jawaban yang benar, melainkan diminta menemukan sendiri pola, "
    "struktur, atau hubungan yang tersembunyi di dalam data. Data yang memiliki karakteristik "
    "mirip akan dikelompokkan menjadi satu sehingga lebih mudah dipahami.")
par(doc, "Tujuan utama unsupervised learning bukan untuk melakukan prediksi, melainkan untuk "
    "memahami karakteristik data. Oleh karena itu, metode ini sering dipakai pada tahap awal "
    "analisis, misalnya untuk mengelompokkan pelanggan berdasarkan kebiasaan belanja, atau "
    "mengelompokkan daerah berdasarkan kondisi sosial-ekonominya. Karena tidak memerlukan proses "
    "pelabelan terlebih dahulu, metode ini sangat menghemat waktu ketika jumlah data sangat banyak.")

judul(doc, "B. Klustering dan K-Means", 2)
par(doc, "Klustering adalah proses mengelompokkan data yang memiliki kemiripan ke dalam satu "
    "kelompok, sedangkan data yang berbeda ditempatkan pada kelompok lain. Teknik ini tidak "
    "membutuhkan label sejak awal sehingga cocok untuk data yang masih acak. Semakin baik "
    "pengelompokannya, semakin jelas pula pola yang terlihat pada data.")
par(doc, "K-Means adalah salah satu algoritma klustering yang paling sering digunakan karena "
    "sederhana dan cepat. Huruf K menunjukkan jumlah kelompok yang ingin dibentuk. Pada awal "
    "proses, sistem memilih beberapa titik pusat (centroid) secara acak. Setiap data kemudian "
    "dihitung jaraknya ke tiap centroid dan dimasukkan ke kelompok dengan centroid terdekat. "
    "Setelah semua data terbagi, posisi centroid dihitung ulang berdasarkan rata-rata anggota "
    "kelompoknya. Proses ini diulang sampai posisi centroid stabil (tidak berubah lagi).")

judul(doc, "C. Centroid", 2)
par(doc, "Centroid adalah titik pusat yang mewakili sebuah kelompok pada algoritma K-Means. "
    "Titik ini menjadi acuan untuk menentukan sebuah data masuk ke kelompok yang mana. Pada "
    "setiap iterasi, posisi centroid diperbarui dengan menghitung rata-rata seluruh data pada "
    "kelompoknya, sehingga posisinya benar-benar berada di tengah kelompok. Ketika posisi "
    "centroid sudah tidak berubah, proses dianggap selesai. Semakin tepat posisi centroid, "
    "semakin baik pula hasil pengelompokan data. Oleh karena itu, sebelum menjalankan K-Means "
    "biasanya dilakukan pembersihan data dan normalisasi agar setiap variabel memiliki skala "
    "yang seimbang.")

# ============================================================
judul(doc, "2. LANGKAH-LANGKAH PEMBUATAN PROGRAM", 1)

judul(doc, "A. Menyiapkan Dataset", 2)
par(doc, "Dataset yang digunakan pada praktikum ini berisi data kondisi sosial-ekonomi dari "
    "205 kabupaten/kota di Indonesia. Data bersumber dari Badan Pusat Statistik (BPS), yaitu "
    "empat komponen penyusun Indeks Pembangunan Manusia (IPM) Metode Baru per kabupaten/kota: "
    "Umur Harapan Hidup (UHH), Harapan Lama Sekolah (HLS), Rata-rata Lama Sekolah (RLS), dan "
    "Pengeluaran per Kapita Disesuaikan. Keempat variabel tersebut mencerminkan aspek kesehatan, "
    "pendidikan, dan ekonomi suatu daerah, sehingga cocok digunakan untuk mengelompokkan tingkat "
    "kerentanan sosial-ekonomi antar wilayah.")
par(doc, "Dataset disiapkan dalam dua versi: versi belum di-encode (kategori kerentanan berupa "
    "teks: Rentan, Berkembang, Sejahtera) dan versi sudah di-encode (kategori berupa angka: 0, "
    "1, 2). Versi angka inilah yang digunakan sebagai target pada proses klasifikasi.")
sisip(doc, "Sisipkan screenshot dataset (Dataset_Awal_TanpaTarget.xlsx) di sini - Gambar 4.1")

judul(doc, "B. Instalasi dan Memuat Data", 2)
par(doc, "Program dibangun menggunakan bahasa Python dengan bantuan pustaka pandas, "
    "scikit-learn, imbalanced-learn (SMOTE), dan matplotlib. Tahap pertama adalah memuat file "
    "Excel, memilih empat kolom numerik, lalu menstandarkan data menggunakan StandardScaler agar "
    "seluruh variabel memiliki skala yang setara. Selanjutnya algoritma K-Means diterapkan "
    "dengan jumlah kluster sebanyak 3.")
sisip(doc, "Sisipkan screenshot Source Code Proses Klustering K-Means - Gambar 4.2")
par(doc, "Hasil klustering divisualisasikan dalam bentuk scatter plot dengan sumbu Rata-rata "
    "Lama Sekolah dan Pengeluaran per Kapita (dalam bentuk standarisasi). Warna yang berbeda "
    "menunjukkan kluster yang berbeda. Dataset yang telah diberi label kluster kemudian disimpan "
    "ke dalam dua file (versi teks dan versi angka).")
sisip(doc, "Sisipkan screenshot grafik hasil klustering (hasil_klustering.png) - Gambar 4.3")
par(doc, "Setelah klustering selesai, data dibaca kembali lalu dipisahkan menjadi variabel "
    "input (X) berisi empat kolom numerik, dan label (Y) berisi kolom kluster. Data dibagi "
    "menjadi data latih dan data uji dengan perbandingan 80% : 20%. Metode SMOTE digunakan untuk "
    "menyeimbangkan jumlah data tiap kelas, dan StandardScaler kembali diterapkan sebelum "
    "pelatihan model klasifikasi Naive Bayes, Random Forest, dan Support Vector Machine (SVM).")
sisip(doc, "Sisipkan screenshot Source Code Klasifikasi (Naive Bayes, Random Forest, SVM) - Gambar 4.4")

# ============================================================
judul(doc, "3. HASIL DAN PENGAMATAN", 1)

judul(doc, "A. Jumlah Kluster dan Hasil Klustering", 2)
par(doc, "Berdasarkan hasil klustering menggunakan algoritma K-Means, 205 kabupaten/kota "
    "berhasil dikelompokkan menjadi 3 kluster berdasarkan kemiripan kondisi sosial-ekonominya. "
    "Setelah nomor kluster diurutkan berdasarkan tingkat kesejahteraan, diperoleh interpretasi "
    "sebagai berikut:")
tabel = doc.add_table(rows=4, cols=3)
tabel.style = "Light Grid Accent 1"
hdr = tabel.rows[0].cells
hdr[0].text = "Kluster"
hdr[1].text = "Kategori"
hdr[2].text = "Jumlah Daerah"
for i, (a, b, c) in enumerate([("0", "Rentan", "102"), ("1", "Berkembang", "66"),
                               ("2", "Sejahtera", "37")], start=1):
    sel = tabel.rows[i].cells
    sel[0].text, sel[1].text, sel[2].text = a, b, c
par(doc, "Kluster Rentan berisi daerah dengan angka harapan hidup, lama sekolah, dan "
    "pengeluaran per kapita yang relatif rendah; kluster Berkembang berada pada tingkat "
    "menengah; sedangkan kluster Sejahtera memuat daerah dengan indikator sosial-ekonomi paling "
    "tinggi (umumnya kota-kota besar). Label kluster inilah yang kemudian digunakan sebagai "
    "target pada proses klasifikasi.")
sisip(doc, "Sisipkan screenshot grafik & tabel hasil klustering - Gambar 4.5")

judul(doc, "B. Perbedaan Akurasi Antara Naive Bayes, SVM, dan Random Forest", 2)
par(doc, "Setelah label kluster terbentuk, dilakukan klasifikasi menggunakan tiga algoritma "
    "untuk memprediksi kategori kerentanan sosial-ekonomi suatu daerah. Pengujian dilakukan "
    "terhadap 41 data uji (20% dari total data). Hasil akurasi ketiga algoritma dirangkum pada "
    "tabel berikut:")
tabel2 = doc.add_table(rows=4, cols=3)
tabel2.style = "Light Grid Accent 1"
h2 = tabel2.rows[0].cells
h2[0].text = "Algoritma"
h2[1].text = "Akurasi"
h2[2].text = "Keterangan"
for i, (a, b, c) in enumerate([("Naive Bayes", "95,12%", "39 benar, 2 salah"),
                               ("Support Vector Machine (SVM)", "95,12%", "39 benar, 2 salah"),
                               ("Random Forest", "95,12%", "39 benar, 2 salah")], start=1):
    sel = tabel2.rows[i].cells
    sel[0].text, sel[1].text, sel[2].text = a, b, c
par(doc, "Berdasarkan hasil pengujian, ketiga algoritma memberikan tingkat akurasi yang sama, "
    "yaitu 95,12%, dengan 39 dari 41 data uji berhasil diprediksi dengan benar. Hal ini "
    "menunjukkan bahwa kluster yang terbentuk dari K-Means cukup terpisah dengan baik sehingga "
    "mudah dikenali oleh berbagai model klasifikasi. Meskipun akurasinya sama, ketiga algoritma "
    "memiliki karakteristik berbeda: Naive Bayes unggul dari sisi kesederhanaan dan kecepatan, "
    "SVM baik dalam memisahkan kelas dengan margin yang jelas, sedangkan Random Forest tahan "
    "terhadap data kompleks karena menggabungkan banyak pohon keputusan.")
sisip(doc, "Sisipkan screenshot Classification Report Naive Bayes, Random Forest, dan SVM - Gambar 4.6")

# ============================================================
judul(doc, "4. KESIMPULAN", 1)
par(doc, "Berdasarkan hasil praktikum, algoritma K-Means berhasil mengelompokkan 205 "
    "kabupaten/kota di Indonesia menjadi 3 kluster tingkat kerentanan sosial-ekonomi, yaitu "
    "Rentan (102 daerah), Berkembang (66 daerah), dan Sejahtera (37 daerah), berdasarkan empat "
    "komponen IPM (UHH, HLS, RLS, dan Pengeluaran per Kapita) dari data BPS. Label kluster "
    "tersebut kemudian digunakan sebagai target pada proses klasifikasi menggunakan Naive Bayes, "
    "Random Forest, dan Support Vector Machine (SVM). Ketiga algoritma memperoleh akurasi yang "
    "sama sebesar 95,12%, yang menandakan hasil klustering sudah cukup baik dan konsisten. "
    "Praktikum ini membantu memahami bagaimana metode unsupervised learning (klustering) dapat "
    "dipadukan dengan metode klasifikasi untuk menganalisis dan memetakan kondisi sosial-ekonomi "
    "daerah.")

par(doc, "")
p = doc.add_paragraph()
r = p.add_run("Sumber data: Badan Pusat Statistik (BPS), www.bps.go.id - Komponen Indeks "
              "Pembangunan Manusia Metode Baru menurut Kabupaten/Kota (UHH 2024; HLS, RLS, "
              "Pengeluaran per Kapita 2025).")
r.italic = True
r.font.size = Pt(9)

doc.save("Laporan_Praktikum4.docx")
print("Laporan disimpan: Laporan_Praktikum4.docx")
